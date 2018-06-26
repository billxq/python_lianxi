#!usr/bin/env python  
# coding:utf-8
""" 
@author:xuqing 
@file: shuxueti.py 
@time: 2018/06/26 
"""


import random
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def get_randint(a,b):
    return random.randint(a,b)

def get_mp():
    l1 = ['+','-']
    return random.sample(l1,1)[0]

def get_q():
    m = get_mp()
    n1 = get_randint(1, 10)
    n2 = get_randint(1, 10)
    if m == '-' and n1<n2:
        m = "+"
        q = str(n1) + '  ' + m + '  ' + str(n2) + ' = '
        return q
    else:
        q = str(n1) + '  ' + m + '  ' + str(n2) + ' = '
        return q

if __name__ == '__main__':
    document = Document()
    paragraph_format = WD_PARAGRAPH_ALIGNMENT
    document.add_heading('小练习',0)
    p = document.add_paragraph('小练习')
    pf = p.paragraph_format
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i in range(30):
        for j in range(3):
            document.add_paragraph(get_q()+'\t'*3 + get_q()+'\t'*3 + get_q()+'\t'*3)
            print(get_q(),end='\t'*3)
        print('\n')

    document.save('demon2.docx')
